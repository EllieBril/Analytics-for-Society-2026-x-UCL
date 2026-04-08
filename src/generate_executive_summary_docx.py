from __future__ import annotations

from datetime import date
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parent.parent
OUTPUTS = ROOT / "outputs"


def add_bullet(document: Document, text: str) -> None:
    document.add_paragraph(text, style="List Bullet")


def set_document_defaults(document: Document) -> None:
    style = document.styles["Normal"]
    style.font.name = "Aptos"
    style.font.size = Pt(10.5)
    for section in document.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)


def add_table(document: Document, dataframe: pd.DataFrame, title: str) -> None:
    document.add_paragraph(title, style="Heading 2")
    table = document.add_table(rows=1, cols=len(dataframe.columns))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for idx, col in enumerate(dataframe.columns):
        hdr_cells[idx].text = str(col)
    for _, row in dataframe.iterrows():
        cells = table.add_row().cells
        for idx, value in enumerate(row.tolist()):
            cells[idx].text = str(value)
    document.add_paragraph("")


def build_summary_inputs() -> dict[str, object]:
    latest = pd.read_csv(OUTPUTS / "03_country_trajectories" / "country_rankings_latest.csv")
    trend = pd.read_csv(OUTPUTS / "03_country_trajectories" / "country_rankings_trend.csv")
    segments = pd.read_csv(OUTPUTS / "08_segmentation" / "segment_profiles.csv")
    correlates = pd.read_csv(OUTPUTS / "06_correlates" / "candidate_driver_rankings.csv")
    intervention_map = pd.read_csv(OUTPUTS / "09_intervention_mapping" / "segment_to_intervention_map.csv")
    school_quadrants = pd.read_csv(OUTPUTS / "07_within_school_gap" / "quadrant_profiles.csv")
    inventory = pd.read_csv(OUTPUTS / "01_data_inventory" / "metadata_summary.csv")
    final_report = (OUTPUTS / "12_final" / "FINAL_REPORT.md").read_text(encoding="utf-8")

    worsening = trend.loc[trend["trend_label"] == "Worsening"].copy()

    top_latest = latest.head(5).copy()
    top_latest["latest_gap"] = top_latest["latest_gap"].round(1)
    top_latest = top_latest[["CNT", "latest_year", "latest_gap", "latest_avg_math"]]
    top_latest.columns = ["System", "Latest year", "Latest gap", "Avg maths"]

    top_worsening = worsening.head(5).copy()
    top_worsening["long_run_change"] = top_worsening["long_run_change"].round(1)
    top_worsening["last_gap"] = top_worsening["last_gap"].round(1)
    top_worsening = top_worsening[["CNT", "first_year", "last_year", "long_run_change", "last_gap"]]
    top_worsening.columns = ["System", "First year", "Last year", "Gap change", "Latest gap"]

    seg_table = segments[["official_segment", "segment_size", "school_mean_math", "within_school_gap", "EQUITY_RISK_SCORE"]].copy()
    seg_table["school_mean_math"] = seg_table["school_mean_math"].round(1)
    seg_table["within_school_gap"] = seg_table["within_school_gap"].round(1)
    seg_table["EQUITY_RISK_SCORE"] = seg_table["EQUITY_RISK_SCORE"].round(1)
    seg_table.columns = ["Segment", "Schools", "Mean maths", "Within-school gap", "Risk score"]

    top_perf = correlates.loc[correlates["outcome"] == "school_mean_math"].head(3).copy()
    top_risk = correlates.loc[correlates["outcome"] == "school_risk_score"].head(3).copy()
    top_gap = correlates.loc[correlates["outcome"] == "within_school_gap"].head(3).copy()

    segment_recs = (
        intervention_map.sort_values(["segment", "category_rank"])
        .groupby("segment")["intervention"]
        .apply(lambda s: " | ".join(s.astype(str)))
        .to_dict()
    )

    return {
        "latest": latest,
        "trend": trend,
        "segments": segments,
        "school_quadrants": school_quadrants,
        "inventory": inventory,
        "final_report": final_report,
        "top_latest": top_latest,
        "top_worsening": top_worsening,
        "seg_table": seg_table,
        "top_perf": top_perf,
        "top_risk": top_risk,
        "top_gap": top_gap,
        "segment_recs": segment_recs,
    }


def build_document() -> Document:
    data = build_summary_inputs()
    latest = data["latest"]
    trend = data["trend"]
    segments = data["segments"]
    quadrants = data["school_quadrants"]
    inventory = data["inventory"]

    document = Document()
    set_document_defaults(document)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Executive Summary\nSDG 4 Equity Diagnostic and School Segmentation Prototype")
    run.bold = True
    run.font.size = Pt(18)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(f"Generated on {date.today().isoformat()} from the completed local analysis pipeline.").italic = True

    document.add_paragraph("")

    document.add_paragraph("Purpose", style="Heading 1")
    document.add_paragraph(
        "This project tests whether available PISA-style data are sufficient to support a rigorous SDG 4 competition concept focused on education inequality. "
        "The final product is a hybrid decision-support prototype: cross-country equity diagnostics, school-profile segmentation, and transparent intervention prioritisation."
    )

    document.add_paragraph("Data and methodology", style="Heading 1")
    add_bullet(
        document,
        f"Data coverage: {inventory['path'].nunique()} local input files were inventoried, including derived CSV tables and raw PISA 2022 SAV files for students, schools, teachers, cognitive items, and financial-literacy files.",
    )
    add_bullet(
        document,
        "Trend analysis used all available country-year data. Rich school-level profiling used the 2022 raw files because that is where the strongest school-context detail exists.",
    )
    add_bullet(
        document,
        "Primary equity metric: weighted mathematics score gap between top and bottom ESCS quartiles at country level. School-level inequality used a within-school SES-maths gap and SES slope where sample size permitted.",
    )
    add_bullet(
        document,
        "Methods were deliberately transparent: weighted descriptive summaries, country ranking and trajectory classification, school risk diagnostics, regularised correlates analysis, within-school quadrant analysis, and interpretable rule-based segmentation.",
    )
    add_bullet(
        document,
        "Interpretation standard: observational and correlational. The analysis supports prioritisation and diagnosis, not causal proof."
    )

    document.add_paragraph("Key findings", style="Heading 1")
    add_bullet(
        document,
        f"Highest latest observed SES-maths gaps were Romania ({latest.iloc[0]['latest_gap']:.1f}), Slovakia ({latest.iloc[1]['latest_gap']:.1f}), Israel ({latest.iloc[2]['latest_gap']:.1f}), Czechia ({latest.iloc[3]['latest_gap']:.1f}), and Chinese Taipei ({latest.iloc[4]['latest_gap']:.1f}).",
    )
    worsening = trend.loc[trend["trend_label"] == "Worsening"].copy()
    add_bullet(
        document,
        f"Among systems with strong trend evidence, the largest gap increases were Romania (+{worsening.iloc[0]['long_run_change']:.1f}), Italy (+{worsening.iloc[1]['long_run_change']:.1f}), Czechia (+{worsening.iloc[2]['long_run_change']:.1f}), Finland (+{worsening.iloc[3]['long_run_change']:.1f}), and Israel (+{worsening.iloc[4]['long_run_change']:.1f}).",
    )
    add_bullet(
        document,
        f"The published school typology identifies {segments['official_segment'].nunique()} interpretable segments across {int(segments['segment_size'].sum()):,} schools with complete enough data for segmentation.",
    )
    largest_segment = segments.sort_values("segment_size", ascending=False).iloc[0]
    add_bullet(
        document,
        f"The largest school segment is '{largest_segment['official_segment']}', combining {int(largest_segment['segment_size']):,} schools with mean maths {largest_segment['school_mean_math']:.1f} and within-school gap {largest_segment['within_school_gap']:.1f}.",
    )
    high_quad = quadrants.loc[quadrants["school_quadrant"] == "High performance / High inequality"].iloc[0]
    low_quad = quadrants.loc[quadrants["school_quadrant"] == "Low performance / Low inequality"].iloc[0]
    add_bullet(
        document,
        f"Within-school inequality matters: high-performance/high-inequality schools average maths {high_quad['school_mean_math']:.1f} with gap {high_quad['within_school_gap']:.1f}, while low-performance/low-inequality schools average maths {low_quad['school_mean_math']:.1f} with gap {low_quad['within_school_gap']:.1f}.",
    )

    document.add_paragraph("Correlates and likely drivers", style="Heading 1")
    add_bullet(
        document,
        f"For school mean maths, the strongest observed correlates were school mean ESCS, student maths self-efficacy, and attendance-related measures such as skipping and tardiness.",
    )
    add_bullet(
        document,
        "For the supplied school risk score, lower homework, weaker ICT regulation, more skipping, and lower educational expectations co-moved with higher risk.",
    )
    add_bullet(
        document,
        "For within-school inequality, school SES composition, homework, belonging, climate, and support-staff proxies were among the strongest observed correlates.",
    )
    add_bullet(
        document,
        "These are correlates, not proven causes. The current data can describe patterns associated with gaps, but not identify the definitive reason the gap exists inside a school."
    )

    document.add_paragraph("Recommended final representable", style="Heading 1")
    add_bullet(
        document,
        "The strongest competition representable is a dashboard prototype with interactive visualisations inside it, not a standalone chart deck."
    )
    add_bullet(
        document,
        "Recommended structure: country diagnostic view, school-segmentation view, recommendation engine, and evidence/limits view."
    )
    add_bullet(
        document,
        "This framing is stronger than a generic policy essay and safer than a fully prescriptive individual-school dashboard."
    )

    document.add_paragraph("Limits and safe claims", style="Heading 1")
    add_bullet(document, "Safe claim: the dashboard identifies where observed equity gaps are high or worsening and which school profiles are associated with higher risk.")
    add_bullet(document, "Safe claim: the tool supports evidence-informed prioritisation of intervention categories by profile.")
    add_bullet(document, "Unsafe claim: the tool identifies the exact cause of a school’s gap.")
    add_bullet(document, "Unsafe claim: the recommended intervention is proven to reduce the gap for one named school.")

    document.add_paragraph("Project recommendation", style="Heading 1")
    add_bullet(
        document,
        "Best framing: hybrid risk scoring + segmentation + recommendation logic."
    )
    add_bullet(
        document,
        "Best scope: all years for the country story, 2022 as the core year for school profiling."
    )
    add_bullet(
        document,
        "Most valuable next data: longitudinal outcomes, intervention implementation/outcome data, and harmonised multi-year school-context variables."
    )

    document.add_page_break()

    add_table(document, data["top_latest"], "Table 1. Highest latest observed SES-maths gaps")
    add_table(document, data["top_worsening"], "Table 2. Largest worsening systems with strong trend evidence")
    add_table(document, data["seg_table"], "Table 3. Published school segments")

    document.add_paragraph("Method note", style="Heading 1")
    document.add_paragraph(
        "This executive summary is generated directly from the completed analysis outputs in the local repository. "
        "It summarises the strongest evidence-backed results and the final product recommendation in concise form."
    )

    return document


def main() -> None:
    output_path = OUTPUTS / "12_final" / "EXECUTIVE_SUMMARY.docx"
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document = build_document()
    document.save(output_path)
    print(output_path)


if __name__ == "__main__":
    main()
